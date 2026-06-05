from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/modelo-padrao-producao-questoes-autores.docx")

GREEN = "1E604A"
TEAL = "2A7D68"
PALE_GREEN = "EAF4EF"
PALE_GOLD = "FFF7E5"
GOLD = "B8862D"
INK = "17251F"
MUTED = "5F6F68"
LINE = "C8D7D0"
WHITE = "FFFFFF"
LIGHT_GRAY = "F5F7F6"

CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=100, start=140, bottom=100, end=140):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[index])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=None, color=None, bold=None, italic=None, name="Calibri"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_page_number(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instruction, separate, text, end])
    set_font(run, size=9, color=MUTED)


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    paragraph.add_run(text)
    keep_with_next(paragraph)
    return paragraph


def add_body(doc, text, bold_prefix=None, italic=False, after=6):
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=after)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_font(prefix, bold=True, color=INK)
        remainder = paragraph.add_run(text[len(bold_prefix) :])
        set_font(remainder, color=INK, italic=italic)
    else:
        run = paragraph.add_run(text)
        set_font(run, color=INK, italic=italic)
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(paragraph, after=4)
    run = paragraph.add_run(text)
    set_font(run, color=INK)
    return paragraph


def add_numbered(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(paragraph, after=4)
    run = paragraph.add_run(text)
    set_font(run, color=INK)
    return paragraph


def add_callout(doc, label, text, fill=PALE_GREEN, accent=GREEN):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=0, line=1.15)
    label_run = paragraph.add_run(f"{label}: ")
    set_font(label_run, size=10.5, color=accent, bold=True)
    text_run = paragraph.add_run(text)
    set_font(text_run, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_metadata_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        set_cell_shading(cells[0], PALE_GREEN)
        label_p = cells[0].paragraphs[0]
        set_paragraph_spacing(label_p, after=0, line=1.1)
        label_run = label_p.add_run(label)
        set_font(label_run, size=9.5, color=GREEN, bold=True)
        value_p = cells[1].paragraphs[0]
        set_paragraph_spacing(value_p, after=0, line=1.1)
        value_run = value_p.add_run(value)
        set_font(value_run, size=10, color=INK)
    set_table_geometry(table, [2600, 6760])
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_field(doc, label, value, *, fill=None, long=False):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA])
    cell = table.cell(0, 0)
    if fill:
        set_cell_shading(cell, fill)
    paragraph = cell.paragraphs[0]
    set_paragraph_spacing(paragraph, after=0, line=1.15)
    label_run = paragraph.add_run(f"{label}: ")
    set_font(label_run, size=9.5, color=GREEN, bold=True)
    value_run = paragraph.add_run(value)
    set_font(value_run, size=10.5, color=INK)
    if long:
        set_cell_margins(cell, top=130, start=160, bottom=130, end=160)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.35
    return table


def add_question_marker(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA], indent_dxa=0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, GREEN)
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=0, line=1.0)
    run = paragraph.add_run(text)
    set_font(run, size=10, color=WHITE, bold=True)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def add_question_block(doc, values, blank=False):
    add_question_marker(doc, "INÍCIO DA QUESTÃO")
    add_metadata_table(
        doc,
        [
            ("CÓDIGO INTERNO", values["code"]),
            ("CARREIRA", values["career"]),
            ("DISCIPLINA", values["subject"]),
            ("DIFICULDADE", values["difficulty"]),
            ("TIPO DE ORIGEM", values["source_type"]),
            ("BANCA", values["board"]),
            ("ÓRGÃO", values["orgao"]),
            ("CARGO", values["cargo"]),
            ("ANO", values["year"]),
            ("REFERÊNCIA DA FONTE", values["reference"]),
        ],
    )
    add_field(doc, "ENUNCIADO", values["statement"], fill=LIGHT_GRAY, long=True)
    for label in "ABCDE":
        add_field(doc, f"ALTERNATIVA {label}", values[f"alt_{label.lower()}"], long=True)
    add_field(doc, "GABARITO", values["correct"], fill=PALE_GOLD)
    add_field(
        doc,
        "COMENTÁRIO GERAL",
        values["general_comment"],
        fill=LIGHT_GRAY,
        long=True,
    )
    for label in "ABCDE":
        add_field(
            doc,
            f"COMENTÁRIO DA ALTERNATIVA {label}",
            values[f"comment_{label.lower()}"],
            long=True,
        )
    if blank:
        add_field(
            doc,
            "OBSERVAÇÃO EDITORIAL",
            "[Opcional. Registre dúvida, imagem necessária ou ponto a revisar.]",
            fill=PALE_GOLD,
            long=True,
        )
    add_question_marker(doc, "FIM DA QUESTÃO")


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    style_tokens = {
        "Heading 1": (16, 18, 10, GREEN),
        "Heading 2": (13, 14, 7, GREEN),
        "Heading 3": (12, 10, 5, TEAL),
    }
    for name, (size, before, after, color) in style_tokens.items():
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25


def configure_page(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.78)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.38)
    section.footer_distance = Inches(0.38)

    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("APROVAENF  |  PRODUÇÃO EDITORIAL")
    set_font(run, size=8.5, color=MUTED, bold=True)

    footer = section.footer
    table = footer.add_table(rows=1, cols=2, width=Inches(6.5))
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_geometry(table, [7200, 2160], indent_dxa=0)
    left = table.cell(0, 0).paragraphs[0]
    set_paragraph_spacing(left, after=0, line=1.0)
    left_run = left.add_run("Modelo padrão de produção de questões")
    set_font(left_run, size=8.5, color=MUTED)
    right = table.cell(0, 1).paragraphs[0]
    right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph_spacing(right, after=0, line=1.0)
    page_label = right.add_run("Página ")
    set_font(page_label, size=8.5, color=MUTED)
    add_page_number(right)


def add_title_page(doc):
    kicker = doc.add_paragraph()
    set_paragraph_spacing(kicker, before=12, after=8, line=1.0)
    run = kicker.add_run("GUIA E MODELO DE PREENCHIMENTO")
    set_font(run, size=10, color=GOLD, bold=True)

    title = doc.add_paragraph()
    set_paragraph_spacing(title, after=6, line=1.0)
    run = title.add_run("Produção padronizada de questões")
    set_font(run, size=27, color=GREEN, bold=True)

    subtitle = doc.add_paragraph()
    set_paragraph_spacing(subtitle, after=20, line=1.15)
    run = subtitle.add_run(
        "Documento para autores do Aprovaenf, preparado para revisão editorial "
        "e conversão estruturada por inteligência artificial."
    )
    set_font(run, size=13, color=MUTED)

    add_callout(
        doc,
        "Regra central",
        "Uma questão deve corresponder a um bloco completo, iniciado por "
        "“INÍCIO DA QUESTÃO” e encerrado por “FIM DA QUESTÃO”. Carreira, "
        "disciplina, dificuldade e banca podem ficar vazias para classificação "
        "posterior na plataforma.",
    )

    add_heading(doc, "Como usar este modelo", 1)
    for item in (
        "Preencha os dados da questão usando exatamente os rótulos apresentados.",
        "Duplique o bloco em branco para cada nova questão; não reorganize os campos.",
        "Escreva cada alternativa e cada comentário em seu campo correspondente.",
        "Entregue o arquivo em formato DOCX. Não transforme o conteúdo em imagem ou PDF.",
        "Antes do envio, execute a lista de conferência ao final do documento.",
    ):
        add_numbered(doc, item)

    add_heading(doc, "Por que a padronização importa", 2)
    add_body(
        doc,
        "O formato reduz erros de interpretação, evita que carreira, cargo, banca e "
        "disciplina sejam confundidos e permite converter o conteúdo para o sistema "
        "sem alterar enunciados ou alternativas."
    )

    add_heading(doc, "Campos que exigem atenção", 2)
    for item in (
        "CARREIRA/ESPECIALIDADE: campo opcional na ingestão. Quando preenchido, "
        "use somente a carreira principal, como “Enfermeiro(a)” ou “Tecnico em "
        "enfermagem”. O cargo da prova pertence ao campo CARGO.",
        "DISCIPLINA: campo opcional na ingestão. Quando preenchido, use o nome "
        "fornecido pelo catálogo editorial do Aprovaenf.",
        "TIPO DE ORIGEM: use somente “autoral” ou “prova_oficial”.",
        "DIFICULDADE: campo opcional na ingestão. Quando preenchido, use somente "
        "“facil”, “media” ou “dificil”.",
        "GABARITO: informe somente uma letra maiúscula entre A e E.",
        "CÓDIGO INTERNO: deve ser único e não pode ser reutilizado em outra questão.",
    ):
        add_bullet(doc, item)


def add_writing_rules(doc):
    doc.add_page_break()
    add_heading(doc, "Regras de produção", 1)

    add_heading(doc, "Conteúdo da questão", 2)
    for item in (
        "O enunciado deve ser autossuficiente e conter todas as informações necessárias.",
        "Use entre duas e cinco alternativas. Para provas objetivas tradicionais, prefira cinco.",
        "Deve existir exatamente uma alternativa correta.",
        "Evite alternativas parcialmente corretas, ambíguas ou dependentes de interpretação não explicitada.",
        "Mantenha paralelismo de linguagem e extensão razoavelmente equilibrada entre as alternativas.",
        "Não revele o gabarito por pistas de tamanho, concordância, repetição ou vocabulário.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Comentários pedagógicos", 2)
    for item in (
        "O comentário geral deve explicar o raciocínio central e o conteúdo cobrado.",
        "O comentário da alternativa correta deve justificar por que ela está correta.",
        "Cada alternativa incorreta deve ter uma explicação específica do erro.",
        "Evite comentários genéricos como “incorreta porque está errada”.",
        "Quando houver norma, protocolo ou diretriz, indique a referência usada.",
    ):
        add_bullet(doc, item)

    add_heading(doc, "Formatação que favorece a conversão", 2)
    for item in (
        "Use texto normal do Word. Não coloque a questão em caixa de texto, coluna ou imagem.",
        "Não use tabelas para organizar enunciado, alternativas ou comentários.",
        "Não escreva duas alternativas no mesmo parágrafo.",
        "Não altere os nomes dos rótulos nem use abreviações próprias.",
        "Pontuação, aspas e ponto e vírgula são permitidos e devem ser preservados.",
        "Quando uma imagem for indispensável, escreva no enunciado "
        "[IMAGEM OBRIGATÓRIA: nome do arquivo] e envie a imagem separadamente.",
    ):
        add_bullet(doc, item)

    add_callout(
        doc,
        "Questões de prova oficial",
        "O texto original, o gabarito e os dados da prova devem ser preservados. "
        "Revisões pedagógicas podem ser feitas nos comentários, nunca no enunciado "
        "ou nas alternativas oficiais.",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    add_callout(
        doc,
        "Classificação pendente",
        "Quando carreira, disciplina, dificuldade ou banca não forem conhecidas, "
        "deixe o respectivo campo vazio. Não invente classificações; o autor "
        "completará esses dados na plataforma antes da publicação.",
    )

    add_heading(doc, "Convenção dos campos", 2)
    add_metadata_table(
        doc,
        [
            ("CÓDIGO INTERNO", "Identificador único da questão."),
            ("CARREIRA", "Carreira existente no Aprovaenf."),
            ("DISCIPLINA", "Disciplina existente para a carreira."),
            ("TIPO DE ORIGEM", "autoral ou prova_oficial."),
            ("REFERÊNCIA DA FONTE", "Código da prova, URL ou referência bibliográfica."),
        ],
    )


def add_example(doc):
    add_heading(doc, "Exemplo preenchido", 1)
    add_body(
        doc,
        "O exemplo demonstra apenas a estrutura esperada. Não o mantenha no arquivo final "
        "enviado ao Aprovaenf.",
        italic=True,
    )
    example = {
        "code": "AUTOR-SILVA-2026-Q001",
        "career": "Enfermeiro(a)",
        "subject": "Saude Publica e SUS",
        "difficulty": "facil",
        "source_type": "autoral",
        "board": "[deixar vazio]",
        "orgao": "[deixar vazio]",
        "cargo": "[deixar vazio]",
        "year": "2026",
        "reference": "BRASIL. Constituição Federal de 1988, art. 196.",
        "statement": (
            "Qual princípio doutrinário do Sistema Único de Saúde garante que todas "
            "as pessoas tenham direito de acesso às ações e aos serviços de saúde?"
        ),
        "alt_a": "Universalidade.",
        "alt_b": "Regionalização.",
        "alt_c": "Hierarquização.",
        "alt_d": "Descentralização.",
        "alt_e": "Participação social.",
        "correct": "A",
        "general_comment": (
            "A universalidade estabelece que a saúde é direito de todas as pessoas e "
            "dever do Estado, sem discriminação no acesso."
        ),
        "comment_a": "Correta. Universalidade assegura acesso à saúde para toda a população.",
        "comment_b": "Incorreta. Regionalização organiza os serviços por regiões de saúde.",
        "comment_c": "Incorreta. Hierarquização ordena os serviços por níveis de complexidade.",
        "comment_d": "Incorreta. Descentralização distribui responsabilidades entre os entes federativos.",
        "comment_e": "Incorreta. Participação social envolve Conselhos e Conferências de Saúde.",
    }
    add_question_block(doc, example)


def add_blank_template(doc):
    doc.add_page_break()
    add_heading(doc, "Modelo para preenchimento", 1)
    add_callout(
        doc,
        "Instrução",
        "Duplique todo o bloco abaixo, do início ao fim, para cada questão. "
        "Substitua os textos entre colchetes e não apague os rótulos.",
    )
    blank = {
        "code": "[identificador único]",
        "career": "[nome exato da carreira]",
        "subject": "[nome exato da disciplina]",
        "difficulty": "[facil | media | dificil]",
        "source_type": "[autoral | prova_oficial]",
        "board": "[banca ou deixar vazio]",
        "orgao": "[órgão ou deixar vazio]",
        "cargo": "[cargo ou deixar vazio]",
        "year": "[AAAA ou deixar vazio]",
        "reference": "[código da prova, URL ou referência bibliográfica]",
        "statement": "[Escreva aqui o enunciado completo.]",
        "alt_a": "[Escreva a alternativa A.]",
        "alt_b": "[Escreva a alternativa B.]",
        "alt_c": "[Escreva a alternativa C ou deixe vazio.]",
        "alt_d": "[Escreva a alternativa D ou deixe vazio.]",
        "alt_e": "[Escreva a alternativa E ou deixe vazio.]",
        "correct": "[A | B | C | D | E]",
        "general_comment": "[Explique o conteúdo e o raciocínio central da questão.]",
        "comment_a": "[Justifique especificamente a alternativa A.]",
        "comment_b": "[Justifique especificamente a alternativa B.]",
        "comment_c": "[Justifique especificamente a alternativa C ou deixe vazio.]",
        "comment_d": "[Justifique especificamente a alternativa D ou deixe vazio.]",
        "comment_e": "[Justifique especificamente a alternativa E ou deixe vazio.]",
    }
    add_question_block(doc, blank, blank=True)


def add_checklist(doc):
    doc.add_page_break()
    add_heading(doc, "Conferência antes do envio", 1)
    add_body(
        doc,
        "Marque todos os itens. Questões que não atendam aos critérios poderão retornar "
        "para ajuste editorial."
    )
    checks = (
        "Cada questão possui INÍCIO DA QUESTÃO e FIM DA QUESTÃO.",
        "Todos os códigos internos são únicos.",
        "Campos de classificação conhecidos seguem os nomes permitidos; os "
        "desconhecidos foram deixados vazios.",
        "O cargo não foi colocado no campo CARREIRA.",
        "O tipo de origem é autoral ou prova_oficial.",
        "O enunciado está completo e não depende de informação ausente.",
        "Existe somente uma alternativa correta.",
        "A letra do gabarito corresponde a uma alternativa preenchida.",
        "O comentário geral explica o conteúdo cobrado.",
        "Todas as alternativas preenchidas possuem comentário específico.",
        "Questões oficiais preservam integralmente o texto original.",
        "Referências, normas e anos foram conferidos.",
        "Nenhuma questão foi inserida como imagem, caixa de texto ou tabela.",
        "O arquivo final está em formato DOCX.",
    )
    for item in checks:
        paragraph = doc.add_paragraph()
        set_paragraph_spacing(paragraph, after=7, line=1.15)
        marker = paragraph.add_run("☐  ")
        set_font(marker, size=13, color=GREEN)
        text = paragraph.add_run(item)
        set_font(text, size=10.5, color=INK)

    add_callout(
        doc,
        "Nome sugerido do arquivo",
        "autor-sobrenome_lote-001_aaaa-mm-dd.docx",
        fill=PALE_GOLD,
        accent=GOLD,
    )
    add_body(
        doc,
        "Versão do modelo: 1.0 | Aprovaenf | Produção editorial",
        italic=True,
        after=0,
    )


def main():
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    add_title_page(doc)
    add_writing_rules(doc)
    add_example(doc)
    add_blank_template(doc)
    add_checklist(doc)

    core = doc.core_properties
    core.title = "Modelo padrão de produção de questões para autores"
    core.subject = "Produção editorial e conversão estruturada para o Aprovaenf"
    core.author = "Aprovaenf"
    core.keywords = "questões, autores, produção editorial, importação, IA"
    core.comments = "Modelo oficial de preenchimento para autores."

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    main()
